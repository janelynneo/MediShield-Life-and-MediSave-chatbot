"""
Statement analyzer — OCR + PII masking + LLM explanation.

Uses Cloudflare Workers AI for OCR on uploaded files (images + PDFs).
"""

import io
import os
import base64
import json as _json
import zipfile
from pathlib import Path

import requests
from dotenv import load_dotenv
from PyPDF2 import PdfReader

from utils.pii_masker import mask_all

BASE_DIR = Path(__file__).parent.parent
load_dotenv(BASE_DIR / ".env")


# ---------------------------------------------------------------------------
# Secrets helpers
# ---------------------------------------------------------------------------

def _get_secret(name: str) -> str:
    """Get a secret from Streamlit secrets (cloud) or env var (local)."""
    try:
        import streamlit as st
        return st.secrets.get(name, os.environ.get(name, ""))
    except ImportError:
        return os.environ.get(name, "")


# ---------------------------------------------------------------------------
# File text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using PyPDF2."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception:
        return ""


def _extract_images_from_docx(file_bytes: bytes) -> list[tuple[str, bytes]]:
    """
    Extract embedded images from a DOCX file (which is a ZIP archive).
    Returns list of (filename, image_bytes) tuples.
    """
    images = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/") and not name.endswith("/"):
                    img_bytes = zf.read(name)
                    filename = os.path.basename(name)
                    images.append((filename, img_bytes))
    except Exception:
        pass
    return images


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Also extracts embedded images and runs them through OCR.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to process Word documents. "
            "Please install it with: uv add python-docx"
        )

    try:
        document = Document(io.BytesIO(file_bytes))
        parts = []

        # Extract paragraph text
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table text (common in statements/benefits docs)
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    parts.append(" | ".join(row_text))

        # Extract embedded images and run through OCR
        images = _extract_images_from_docx(file_bytes)
        for filename, img_bytes in images:
            try:
                ocr_text = extract_text_from_image(img_bytes)
                if ocr_text.strip():
                    parts.append(f"[Image: {filename}]\n{ocr_text}")
            except Exception as e:
                # OCR failed for this image — skip it silently
                pass

        return "\n".join(parts)
    except Exception:
        return ""


def _get_image_mime_type(file_bytes: bytes) -> str:
    """Detect MIME type from image bytes."""
    if file_bytes.startswith(b"\x89PNG"):
        return "image/png"
    elif file_bytes.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    elif file_bytes.startswith(b"RIFF") and b"WEBP" in file_bytes[:12]:
        return "image/webp"
    elif file_bytes.startswith(b"II\x2a\x00") or file_bytes.startswith(b"MM\x00\x2a"):
        return "image/tiff"
    elif file_bytes.startswith(b"GIF87a") or file_bytes.startswith(b"GIF89a"):
        return "image/gif"
    elif file_bytes.startswith(b"BM"):
        return "image/bmp"
    return "image/jpeg"  # fallback


def _detect_file_type(file_bytes: bytes, filename: str) -> str:
    """
    Detect the actual file type from bytes, falling back to extension.
    Handles misnamed files (e.g., PNG saved as .docx).
    """
    # Check magic bytes first
    if file_bytes.startswith(b"\x89PNG"):
        return "png"
    elif file_bytes.startswith(b"\xff\xd8\xff"):
        return "jpg"
    elif file_bytes.startswith(b"RIFF") and b"WEBP" in file_bytes[:12]:
        return "webp"
    elif file_bytes.startswith(b"II\x2a\x00") or file_bytes.startswith(b"MM\x00\x2a"):
        return "tiff"
    elif file_bytes.startswith(b"GIF87a") or file_bytes.startswith(b"GIF89a"):
        return "gif"
    elif file_bytes.startswith(b"BM"):
        return "bmp"
    elif file_bytes.startswith(b"%PDF"):
        return "pdf"
    elif file_bytes.startswith(b"PK\x03\x04"):  # DOCX/DOCX-based formats (ZIP)
        return os.path.splitext(filename)[-1].lower().lstrip(".")
    # Not a recognized format — use extension
    return os.path.splitext(filename)[-1].lower().lstrip(".") or "unknown"


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from an image using OCR.space API.

    Works with: PNG, JPEG, PDF, DOCX, TIFF, WEBP, GIF, BMP
    """
    try:
        api_key = _get_secret("OCR_SPACE_API_KEY")

        if not api_key or api_key == "YOUR_OCR_SPACE_KEY_HERE":
            return ""  # API key not configured — skip OCR silently

        url = "https://api.ocr.space/parse/image"
        mime_type = _get_image_mime_type(file_bytes)
        b64_data = base64.b64encode(file_bytes).decode("utf-8")

        payload = {
            "base64Image": f"data:{mime_type};base64,{b64_data}",
            "language": "eng",
            "isOverlayRequired": False,
            "filetype": mime_type.replace("image/", "").upper(),
        }

        headers = {
            "apikey": api_key,
        }

        response = requests.post(url, headers=headers, data=payload, timeout=60)
        response.raise_for_status()

        result = response.json()
        # OCR.space returns {"ParsedResults": [{"ParsedText": "..."}]}
        if result.get("ParsedResults"):
            return result["ParsedResults"][0].get("ParsedText", "")
        return ""
    except Exception:
        # OCR failed — return empty
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type and extract text.

    Supports: PDF (.pdf), DOCX (.docx), images (.png/.jpg/.jpeg/.webp/.tiff/.bmp/.gif).
    Detects file type from magic bytes, not just filename extension.
    """
    ext = _detect_file_type(file_bytes, filename)

    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        # All other extensions treated as images
        return extract_text_from_image(file_bytes)


# ---------------------------------------------------------------------------
# LLM statement analysis
# ---------------------------------------------------------------------------

def analyze_statement(raw_text: str, masked_text: str) -> str:
    """
    Send masked statement text to LLM for explanation.

    Args:
        raw_text: Original extracted text (for debugging/logging — not sent to LLM).
        masked_text: PII-masked version of the statement text.

    Returns:
        Plain-English explanation of the statement.
    """
    try:
        from utils.llm import get_llm
    except ImportError:
        # Fallback if called outside app context
        from langchain_openai import ChatOpenAI
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)

    prompt = _build_analysis_prompt(masked_text)

    try:
        llm = get_llm()
    except Exception:
        from langchain_openai import ChatOpenAI
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)

    try:
        response = llm.invoke([
            {"role": "system", "content": "You are a helpful healthcare assistant specializing in Singapore's MediShield Life and MediSave scheme."},
            {"role": "user", "content": prompt},
        ])
        return response.content
    except Exception as e:
        return (
            "⚠️ Sorry, something went wrong while analyzing your statement. "
            "Please try again or ensure your file is clearly readable."
        )


def _build_analysis_prompt(masked_text: str) -> str:
    return f"""You are analyzing a masked MediShield Life or MediSave statement.
All personal identifying information has been replaced with [MASKED] labels.

Below is the statement text:
---
{masked_text}
---

Please explain this statement in plain English. Structure your response to cover:

1. **Summary** — What type of document is this (e.g., claim statement, hospital bill, benefit summary)?
2. **What was claimed / charged** — Total bill amount and what it covers
3. **Claim payout breakdown** — How much MediShield Life paid, how much MediSave was used, and any patient responsibility
4. **How the payout was calculated** — Explain if deductible, co-insurance, or pro-ration applied
5. **Shortfalls** — Any amounts not covered by MediShield Life or MediSave, and why
6. **What the patient needs to pay** — Net amount payable by the patient

If certain information is not present in the statement, say so clearly.
Keep your explanation clear and easy to understand for a non-expert.
Format amounts as $X,XXX.
If you cannot determine a value, say "not available in this statement" rather than guessing."""
